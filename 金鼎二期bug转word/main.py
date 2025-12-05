import csv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT


def csv_to_word_bug_report(csv_path, output_path):
    # 读取CSV文件
    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        headers = [header.strip('"') for header in next(reader)]  # 读取并清洗标题行
        bugs = list(reader)

    # 创建Word文档
    doc = Document()

    # 设置全局文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)

    # 处理每个Bug
    for bug in bugs:
        # 创建两列表格（字段名 + 字段值）
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'  # 使用预定义的表格样式

        # 设置表格对齐方式和列宽
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Inches(1.5)  # 字段名称列宽
        table.columns[1].width = Inches(4.5)  # 字段值列宽

        # 设置表头
        # hdr_row = table.rows[0]
        # hdr_row.cells[0].text = "字段名称"
        # hdr_row.cells[1].text = "字段内容"
        # hdr_row.cells[0].paragraphs[0].runs[0].bold = True
        # hdr_row.cells[1].paragraphs[0].runs[0].bold = True

        # 填充数据
        for header, value in zip(headers, bug):
            if header in ['所属模块', ' 所属迭代', '相关需求', '相关任务', '']:
                continue
            new_row = table.add_row()

            # 处理字段名称单元格
            name_cell = new_row.cells[0]
            name_cell.text = header
            name_cell.paragraphs[0].runs[0].font.size = Pt(10)

            # 处理字段值单元格
            value_cell = new_row.cells[1]
            cleaned_value = value.strip('"').replace('\r\n', '\n')  # 清洗数据
            value_cell.text = cleaned_value if cleaned_value else "无"
            value_cell.paragraphs[0].runs[0].font.size = Pt(10)

        # 在两个表格之间添加分页符
        doc.add_paragraph()

    # 保存文档
    doc.save(output_path)


# 使用示例（根据实际路径修改）
csv_to_word_bug_report(r'C:\Users\Administrator\Downloads\接口bug.csv', './接口Bug报告.docx')
